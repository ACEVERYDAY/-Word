# *-* encoding:utf-8 *-*
import os
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor


to_path = r'D:\python_project\TableAfterProcessing'
dir_path = r'D:\python_project\XXXX\backup'
name_list = []
for i in os.listdir(dir_path):
    name_list.append(i)

print(name_list)
# ['D:\\python_project\\操作岗位理论考试题库\\backup\\鼎盛成-产品发运工.csv']

for file_name in name_list:
    # 文件读取路径
    from_path = os.path.join(dir_path, file_name)
    # 创建文档对象，设置字体
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    print((file_name))
    p_data = pd.read_csv(from_path, engine='python', usecols=['题型', '题干', '选项', '答案'])
    p_data = p_data.where(p_data.notnull(), '')
    # 读取指定几列
    q_type = ''
    for index in range(len(p_data)):
        r_type = p_data['题型'][index]
        r_cont = p_data['题干'][index]
        r_choose = p_data['选项'][index]
        r_ans = p_data['答案'][index]
        # print(str(r_choose))
        # 判断当前题型，确定是否创建对应类别标题
        if r_type is not q_type:
            q_type = r_type
            document.add_heading(q_type)
        # 将题号以及题干写入文档
        p = document.add_paragraph(str(index+1) + r'.' + str(r_cont))
        # 写入选项
        # if str(r_choose).strip() is not '':
        #     document.add_paragraph(str(r_choose))
        # 若为判断题则将答案写入
        if q_type == '判断题':
            document.add_paragraph(u'答案：' + str(r_ans) + '\n')
        # 否则只标红正确选项
        else:
            res_list = (str(r_choose)).split()
            # print(res_list)
            p.add_run('\n')
            for res in res_list:
                run = p.add_run(str(res) + ' ')
                # print(res[0])
                if res[0] in r_ans:
                    run.font.color.rgb = RGBColor(255,0,0)
            p.add_run('\n')
            # p.add_run('')
            # 切分答案字符串
        # pass
    # 写入对应路径
    document.save(os.path.join(to_path, file_name[0:-4]+'.docx'))



